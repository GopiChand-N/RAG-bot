import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin
from PyPDF2 import PdfReader
from io import BytesIO
from docx import Document
import json
from html2text import HTML2Text  # Import required for HTML to markdown conversion

def get_data_from_website(url):
    from utils import is_url_already_stored, normalize_url
    """
    Fetches and processes data from a website. Handles directly linked PDFs or documents,
    HTML content, and embedded or downloadable files.

    Args:
        url (str): The URL of the website.

    Returns:
        dict: A dictionary containing the processed content and metadata.
    """
    try:
        response = requests.get(url, timeout=1000)
        if response.status_code != 200:
            print(f"Error fetching {url}: {response.status_code}")
            return {"error": f"Failed to fetch URL. Status code: {response.status_code}"}

        # Determine the content type
        content_type = response.headers.get('Content-Type', '').lower()

        # Handle directly linked PDFs
        if 'application/pdf' in content_type:
            print(f"Processing PDF from URL: {url}")
            pdf_reader = PdfReader(BytesIO(response.content))
            pdf_text = "\n".join(page.extract_text() for page in pdf_reader.pages)
            return {"content": pdf_text, "metadata": {"url": normalize_url(url), "type": "PDF"}}

        # Handle directly linked Word documents
        if 'application/msword' in content_type or 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' in content_type:
            print(f"Processing Word document from URL: {url}")
            doc = Document(BytesIO(response.content))
            doc_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return {"content": doc_text, "metadata": {"url": normalize_url(url), "type": "Word Document"}}

        # Handle HTML content
        print(f"Processing HTML content from URL: {url}")
        soup = BeautifulSoup(response.content, "html.parser")

        # Remove JavaScript and CSS code
        for script in soup(["script", "style"]):
            script.extract()

        # Convert HTML to markdown text
        html2text_instance = HTML2Text()
        html2text_instance.images_to_alt = True
        html2text_instance.body_width = 0
        html2text_instance.single_line_break = True
        text = html2text_instance.handle(str(soup))

        # Extract metadata
        title = soup.title.string.strip() if soup.title else url.split("/")[-1] or "Untitled"
        meta_description = soup.find("meta", attrs={"name": "description"})
        meta_keywords = soup.find("meta", attrs={"name": "keywords"})

        metadata = {
            "title": title,
            "url": normalize_url(url),
            "description": meta_description.get("content").strip() if meta_description else title,
            "keywords": meta_keywords.get("content").strip() if meta_keywords else "",
            "type": "HTML",
        }

        # Find links and downloadable files
        links = []
        for a in soup.find_all('a', href=True):
            full_url = urljoin(url, a['href'])
            links.append(full_url)

        pdf_links = []
        for link in links:
            if link.lower().endswith('.pdf'):
                pdf_links.append(link)

        doc_links = []
        for link in links:
            if link.lower().endswith(('.docx', '.doc')):
                doc_links.append(link)

        # Look for downloadable options in buttons or attributes
        for button in soup.find_all(['a', 'button']):
            download_url = button.get('data-download-url')
            if download_url:
                full_download_url = urljoin(url, download_url)
                if full_download_url.lower().endswith('.pdf'):
                    pdf_links.append(full_download_url)
                elif full_download_url.lower().endswith(('.docx', '.doc')):
                    doc_links.append(full_download_url)

        linked_content = []

        # Parse linked PDFs
        for pdf_link in pdf_links:
            if is_url_already_stored(pdf_link):
                print(f"PDF URL already stored: {pdf_link}")
                continue

            print(f"Processing linked PDF: {pdf_link}")
            pdf_response = requests.get(pdf_link, timeout=1000)
            if pdf_response.status_code == 200:
                pdf_reader = PdfReader(BytesIO(pdf_response.content))
                pdf_text = "\n".join(page.extract_text() for page in pdf_reader.pages)
                linked_content.append({"url": pdf_link, "type": "PDF", "content": pdf_text})

        # Parse linked Word documents
        for doc_link in doc_links:
            if is_url_already_stored(doc_link):
                print(f"Word document URL already stored: {doc_link}")
                continue
            print(f"Processing linked Word document: {doc_link}")
            doc_response = requests.get(doc_link, timeout=1000)
            if doc_response.status_code == 200:
                doc = Document(BytesIO(doc_response.content))
                doc_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                linked_content.append({"url": doc_link, "type": "Word Document", "content": doc_text})

        # Return content and metadata for main HTML content and linked content
        metadata["linked_content"] = json.dumps(linked_content)
        return {"content": text, "metadata": metadata}

    except Exception as e:
        print(f"Error processing {url}: {e}")
        return {"error": f"An error occurred: {str(e)}"}


def fetch_single_url_content(url):
    """
    Fetches and processes the content of a single URL by calling the `get_data_from_website` function.

    Args:
        url (str): The URL to fetch.

    Returns:
        dict: A dictionary containing the processed content and metadata.
    """
    try:
        print(f"Fetching content for URL: {url}")
        result = get_data_from_website(url)
        if "error" in result:
            print(f"Error fetching data for URL: {url}")
            return {"error": f"Failed to process URL: {url}", "details": result["error"]}
        return result

    except Exception as e:
        print(f"Error in fetch_single_url_content: {e}")
        return {"error": f"An exception occurred while processing URL: {url}", "details": str(e)}


def fetch_all_links_content(main_url):
    """
    Fetch content from all links found on the main URL, count the number of links,
    and return fetched and skipped links.

    Args:
        main_url (str): The main URL to fetch links from.

    Returns:
        dict: A dictionary containing fetched content, link count, fetched links, and skipped links.
    """
    try:
        print(f"Attempting to fetch content from URL: {main_url}")
        response = requests.get(main_url, timeout=1000)
        if response.status_code != 200:
            print(f"Error fetching {main_url}: Status code {response.status_code}")
            return {"content": [], "count": 0, "fetched_links": [], "skipped_links": []}

        # Parse HTML and extract links
        soup = BeautifulSoup(response.content, 'html.parser')
        links = [urljoin(main_url, a['href']) for a in soup.find_all('a', href=True)]

        # Deduplicate links
        unique_links = list(set(links))
        print(f"Found {len(unique_links)} unique links.")

        all_content = []
        fetched_links = []
        skipped_links = []

        for link in unique_links:
            try:
                print(f"Fetching data for link: {link}")
                result = get_data_from_website(link)
                if "error" not in result and "content" in result and "metadata" in result:
                    print(f"Fetched data for {link}")
                    all_content.append(result)
                    fetched_links.append(link)
                else:
                    print(f"Skipped link {link} due to fetch error.")
                    skipped_links.append(link)
            except Exception as fetch_error:
                print(f"Error processing link {link}: {fetch_error}")
                skipped_links.append(link)

        return {
            "content": all_content,
            "count": len(unique_links),
            "fetched_links": fetched_links,
            "skipped_links": skipped_links
        }

    except Exception as e:
        print(f"Error fetching links from {main_url}: {e}")
        return {"content": [], "count": 0, "fetched_links": [], "skipped_links": []}

def monitor_website(main_url, check_interval=300):
    """
    Monitor a website for changes at specified intervals.

    Args:
        main_url (str): The main URL to monitor.
        check_interval (int): Time in seconds between checks.
    """
    import time

    previous_content = fetch_all_links_content(main_url)

    while True:
        time.sleep(check_interval)
        current_content = fetch_all_links_content(main_url)

        if current_content != previous_content:
            print("Website content has changed.")
            for text, metadata in current_content:
                print(f"Updated Content from {metadata['url']}:")
                print(text)
            previous_content = current_content
        else:
            print("No changes detected.")

# Example usage
# if __name__ == "__main__":
    # main_url = "https://www.infojiniconsulting.com/"  # Replace with your target URL
    # result =fetch_all_links_content(main_url)  # Monitor the site every 5 minutes
    # print("///////////////////////////", jsonify(result))